import ast
import asyncio
import inspect
import json
import logging
import math
import os
import shutil
import subprocess
import sys
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Optional, Union

logger = logging.getLogger(__name__)

# Centralised output directory — must be inside workspace/ so the
# workspace file API can serve files without a path-mismatch.
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
WORKSPACE_OUTPUT_DIR = str(_PROJECT_ROOT / "workspace" / "output")
os.makedirs(WORKSPACE_OUTPUT_DIR, exist_ok=True)


@dataclass
class ToolDefinition:
    """Definition and execution metadata for an agent tool."""
    name: str
    description: str
    parameters: dict[str, Any]
    func: Callable[..., Any]
    required_skill: str = ""

    def to_ollama_tool(self) -> dict[str, Any]:
        """Convert to standard Ollama / OpenAI tool schema."""
        return {
            "type": "function",
            "function": {
                "name": self.name,
                "description": self.description,
                "parameters": self.parameters,
            },
        }



# 1. Safe Mathematical Expression Evaluator


SAFE_MATH_NAMES = {
    "abs": abs,
    "round": round,
    "min": min,
    "max": max,
    "sum": sum,
    "pow": pow,
    "sqrt": math.sqrt,
    "sin": math.sin,
    "cos": math.cos,
    "tan": math.tan,
    "log": math.log,
    "log10": math.log10,
    "log2": math.log2,
    "exp": math.exp,
    "ceil": math.ceil,
    "floor": math.floor,
    "pi": math.pi,
    "e": math.e,
}


class _SafeMathVisitor(ast.NodeVisitor):
    """AST visitor that only evaluates safe mathematical operations."""

    def visit(self, node: ast.AST) -> Any:
        if isinstance(node, ast.Expression):
            return self.visit(node.body)
        elif isinstance(node, ast.Constant):
            if isinstance(node.value, (int, float)):
                return node.value
            raise ValueError(f"Unsupported constant type: {type(node.value)}")
        elif isinstance(node, ast.Name):
            if node.id in SAFE_MATH_NAMES:
                return SAFE_MATH_NAMES[node.id]
            raise ValueError(f"Unknown variable or function: '{node.id}'")
        elif isinstance(node, ast.BinOp):
            left = self.visit(node.left)
            right = self.visit(node.right)
            if isinstance(node.op, ast.Add):
                return left + right
            elif isinstance(node.op, ast.Sub):
                return left - right
            elif isinstance(node.op, ast.Mult):
                return left * right
            elif isinstance(node.op, ast.Div):
                return left / right
            elif isinstance(node.op, ast.FloorDiv):
                return left // right
            elif isinstance(node.op, ast.Mod):
                return left % right
            elif isinstance(node.op, ast.Pow):
                if right > 1000:
                    raise ValueError("Exponent too large")
                return left ** right
            raise ValueError(f"Unsupported operator: {type(node.op)}")
        elif isinstance(node, ast.UnaryOp):
            operand = self.visit(node.operand)
            if isinstance(node.op, ast.UAdd):
                return +operand
            elif isinstance(node.op, ast.USub):
                return -operand
            raise ValueError(f"Unsupported unary operator: {type(node.op)}")
        elif isinstance(node, ast.Call):
            func = self.visit(node.func)
            if not callable(func):
                raise ValueError(f"Object {func} is not callable")
            args = [self.visit(a) for a in node.args]
            return func(*args)
        elif isinstance(node, ast.List):
            return [self.visit(elt) for elt in node.elts]
        elif isinstance(node, ast.Tuple):
            return tuple(self.visit(elt) for elt in node.elts)
        else:
            raise ValueError(f"Unsupported syntax: {type(node).__name__}")


def calculate_expression(expression: str) -> dict[str, Any]:
    """
    Safely evaluate a mathematical expression without eval() vulnerabilities.
    
    Args:
        expression: String mathematical expression (e.g., '14 * 2.5 + sqrt(144)')
    """
    expr_clean = expression.strip()
    try:
        parsed = ast.parse(expr_clean, mode="eval")
        visitor = _SafeMathVisitor()
        result = visitor.visit(parsed)
        return {
            "success": True,
            "expression": expr_clean,
            "result": result,
            "formatted_result": f"{result}",
        }
    except Exception as e:
        return {
            "success": False,
            "expression": expr_clean,
            "error": str(e),
        }


# ==========================================
# 2. Isolated Python Sandbox Runner
# ==========================================

def run_python_sandbox(
    code: str,
    timeout_seconds: int = 5,
    max_output_chars: int = 4000,
    work_dir: Optional[str] = None,
) -> dict[str, Any]:
    """
    Execute Python code in an isolated subprocess sandbox.
    
    Args:
        code: Python script string to execute.
        timeout_seconds: Maximum allowed execution time in seconds.
        max_output_chars: Maximum captured stdout/stderr length.
        work_dir: Optional working directory for execution.
    """
    start_time = time.time()

    # Prepend basic security restraints
    restricted_preamble = (
        "import sys, os\n"
        "# Disable network / socket access in standard library if imported\n"
        "try:\n"
        "    import socket\n"
        "    def _blocked(*args, **kwargs):\n"
        "        raise PermissionError('Network access is disabled in air-gap sandbox')\n"
        "    socket.socket = _blocked\n"
        "    socket.create_connection = _blocked\n"
        "except Exception:\n"
        "    pass\n\n"
    )

    full_script = restricted_preamble + code

    target_dir = work_dir or tempfile.mkdtemp(prefix="vaultmind_sandbox_")
    try:
        script_path = os.path.join(target_dir, "sandbox_run.py")
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(full_script)

        # Isolated environment variables
        env = {
            "PATH": os.environ.get("PATH", "/usr/bin:/bin"),
            "PYTHONPATH": target_dir,
            "PYTHONDONTWRITEBYTECODE": "1",
            "HF_HUB_OFFLINE": "1",
            "TRANSFORMERS_OFFLINE": "1",
        }

        try:
            process = subprocess.Popen(
                [sys.executable, script_path],
                cwd=target_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                env=env,
            )
            stdout, stderr = process.communicate(timeout=timeout_seconds)
            exit_code = process.returncode
            duration = time.time() - start_time

            # Truncate output if necessary
            if len(stdout) > max_output_chars:
                stdout = stdout[:max_output_chars] + "\n...[Output truncated]"
            if len(stderr) > max_output_chars:
                stderr = stderr[:max_output_chars] + "\n...[Error truncated]"

            return {
                "success": exit_code == 0,
                "exit_code": exit_code,
                "stdout": stdout.strip(),
                "stderr": stderr.strip(),
                "execution_time_seconds": round(duration, 3),
            }

        except subprocess.TimeoutExpired:
            process.kill()
            try:
                stdout, stderr = process.communicate(timeout=1)
            except Exception:
                pass
            return {
                "success": False,
                "exit_code": -1,
                "error": f"Execution timed out after {timeout_seconds} seconds.",
                "execution_time_seconds": timeout_seconds,
            }

        except Exception as e:
            return {
                "success": False,
                "exit_code": -1,
                "error": str(e),
                "execution_time_seconds": round(time.time() - start_time, 3),
            }
    finally:
        if work_dir is None and os.path.exists(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)


# ==========================================
# 3. Document Search & Page Reader Tools
# ==========================================

# Mock In-Memory Store for fallback / demonstration
_DOCUMENT_STORE: dict[str, dict[str, Any]] = {
    "doc-safety-std": {
        "title": "Industrial Safety Standards Manual (v4.1)",
        "department": "operations",
        "pages": {
            1: "Industrial Safety Standards Manual - General Overview & Purpose. Covers operating rules and compliance requirements.",
            4: "Section 4.1 Safety Protocol: All automated sorting conveyors must operate under 65 dB noise threshold with redundant emergency stops installed every 15 meters.",
            5: "Section 4.2 Hazardous Material Handling: Class B chemicals require double-walled containment and automated leak sensors.",
        },
    },
    "doc-audit-2026": {
        "title": "Annual Facility Compliance Audit 2026",
        "department": "compliance",
        "pages": {
            1: "Executive Summary: 2026 Facility Inspection Report. Compliance rate reached 98.4%.",
            2: "Detailed Findings: Section 4.1 safety standards were verified across all 12 manufacturing bays on Jan 15, 2026.",
        },
    },
}


def search_documents(
    query: str,
    top_k: int = 3,
    department: Optional[str] = None,
) -> dict[str, Any]:
    """
    Vector similarity search in document database.
    
    Args:
        query: Natural language query string.
        top_k: Maximum number of relevant chunks to retrieve.
        department: Optional filter by department or category.
    """
    # 1. Query Sovereign RAG hybrid retriever if indexed documents exist
    try:
        from backend.rag.service import default_rag_service
        if default_rag_service.count() > 0:
            contexts = default_rag_service.pipeline.retrieve(question=query, n_results=top_k)
            if contexts:
                docs = []
                for ctx in contexts:
                    meta = ctx.get("metadata", {})
                    dist = float(ctx.get("distance", 0.0))
                    sim = max(0.0, round(1.0 - dist, 4))
                    docs.append({
                        "document_id": meta.get("document_id") or meta.get("source", "Document"),
                        "document_title": meta.get("source", "Document"),
                        "page_number": meta.get("page", 1),
                        "department": meta.get("department", "general"),
                        "content": ctx.get("text", ""),
                        "similarity_score": sim,
                    })

                # Also search past conversation history memory across all sessions
                try:
                    import re
                    from .agent_loop import react_agent_engine
                    q_tokens = [t.lower() for t in re.findall(r"\w+", query) if len(t) > 2]
                    matching_chat_turns = []
                    for cid, msgs in getattr(react_agent_engine, "_conversations", {}).items():
                        for m in msgs:
                            txt = m.get("content", "")
                            txt_lower = txt.lower()
                            matches = sum(1 for tok in q_tokens if tok in txt_lower)
                            if matches > 0:
                                matching_chat_turns.append((matches, cid, m.get("role"), txt))

                    matching_chat_turns.sort(key=lambda x: x[0], reverse=True)
                    for _, cid, role, txt in matching_chat_turns[:3]:
                        docs.insert(0, {
                            "document_id": f"chat_history_{cid}",
                            "document_title": f"Previous Conversation ({role})",
                            "page_number": 1,
                            "department": "chat_memory",
                            "content": txt[:1000],
                            "similarity_score": 0.96,
                        })
                except Exception as mem_err:
                    logger.debug(f"Chat memory retrieval: {mem_err}")

                return {
                    "query": query,
                    "results_count": len(docs),
                    "documents": docs[:max(top_k, 5)],
                }
    except Exception as exc:
        logger.debug(f"RAG service query fallback: {exc}")

    # 2. Fallback to built-in reference store for unit tests / offline mock
    from .embedding import cosine_similarity, default_embedder

    query_vec = default_embedder.embed_query(query)
    results = []

    for doc_id, doc in _DOCUMENT_STORE.items():
        if department and doc.get("department", "").lower() != department.lower():
            continue

        for page_num, page_content in doc.get("pages", {}).items():
            page_vec = default_embedder.embed_text(page_content)
            sim = cosine_similarity(query_vec, page_vec)
            results.append({
                "document_id": doc_id,
                "document_title": doc["title"],
                "page_number": page_num,
                "department": doc.get("department", "general"),
                "content": page_content,
                "similarity_score": round(sim, 4),
            })

    # Sort descending by similarity
    results.sort(key=lambda x: x["similarity_score"], reverse=True)
    top_results = results[:top_k]

    return {
        "query": query,
        "results_count": len(top_results),
        "documents": top_results,
    }


def get_document_page(doc_id: str, page_number: int) -> dict[str, Any]:
    """
    Retrieve raw text / markdown for a specific document page.
    
    Args:
        doc_id: Unique identifier of the document.
        page_number: Page number to retrieve (1-indexed).
    """
    # Try RAG vector store by source / metadata
    try:
        from backend.rag.service import default_rag_service
        if default_rag_service.count() > 0:
            chunks = default_rag_service.retriever.store.get_by_source(doc_id)
            if not chunks:
                chunks = default_rag_service.retriever.store.get_by_metadata("document_id", [doc_id])

            matching = [c for c in chunks if c.get("metadata", {}).get("page") == page_number]
            if matching:
                content = "\n\n".join([c.get("text") or c.get("document", "") for c in matching])
                return {
                    "success": True,
                    "document_id": doc_id,
                    "document_title": doc_id,
                    "page_number": page_number,
                    "content": content,
                }
    except Exception as exc:
        logger.debug(f"RAG get_document_page fallback: {exc}")

    doc = _DOCUMENT_STORE.get(doc_id)
    if not doc:
        return {
            "success": False,
            "error": f"Document with ID '{doc_id}' not found.",
        }

    pages = doc.get("pages", {})
    if page_number not in pages:
        return {
            "success": False,
            "document_id": doc_id,
            "document_title": doc.get("title", "Untitled"),
            "error": f"Page {page_number} not found. Document has pages: {list(pages.keys())}",
        }

    return {
        "success": True,
        "document_id": doc_id,
        "document_title": doc.get("title", "Untitled"),
        "page_number": page_number,
        "content": pages[page_number],
    }



# ==========================================
# 4. Report File Generator Tool
# ==========================================

def create_pdf_document(
    title: str,
    content: str,
    output_filename: Optional[str] = None,
    output_dir: str = WORKSPACE_OUTPUT_DIR,
    table_data: Optional[list[list[Any]]] = None,
) -> dict[str, Any]:
    """Generate a professional, fully styled PDF report using ReportLab."""
    os.makedirs(output_dir, exist_ok=True)
    if not output_filename:
        clean_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).rstrip()
        filename_base = clean_title.lower().replace(" ", "_") or "report"
        output_filename = f"{filename_base}_{int(time.time())}.pdf"
    elif not output_filename.endswith(".pdf"):
        output_filename = f"{output_filename}.pdf"

    file_path = os.path.join(output_dir, output_filename)

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors

        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
        )
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=12,
        )
        heading_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#1e293b'),
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceAfter=8,
        )

        elements = []
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 10))

        for line in content.split("\n"):
            clean_line = line.strip()
            if not clean_line:
                elements.append(Spacer(1, 6))
            elif clean_line.startswith("#"):
                header_txt = clean_line.lstrip("#").strip()
                elements.append(Paragraph(header_txt, heading_style))
            elif clean_line.startswith("- ") or clean_line.startswith("* "):
                bullet_txt = f"• {clean_line[2:]}"
                elements.append(Paragraph(bullet_txt, body_style))
            else:
                elements.append(Paragraph(clean_line, body_style))

        if table_data and len(table_data) > 0:
            elements.append(Spacer(1, 12))
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e293b')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            elements.append(t)

        doc.build(elements)
        size_bytes = os.path.getsize(file_path)

        return {
            "success": True,
            "filename": output_filename,
            "file_path": os.path.abspath(file_path),
            "relative_path": f"output/{output_filename}",
            "size_bytes": size_bytes,
            "message": f"Successfully generated PDF: {output_filename}",
        }
    except Exception as e:
        logger.error(f"Failed to generate PDF: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"PDF generation error: {str(e)}",
        }


def resolve_workspace_file_path(file_path: str) -> Optional[Path]:
    """Helper to locate a file across workspace directories."""
    raw_p = Path(file_path.strip())
    if raw_p.is_absolute() and raw_p.exists():
        return raw_p

    clean = file_path.strip().lstrip("/\\")
    direct = Path(clean)
    if direct.is_absolute() and direct.exists():
        return direct

    candidates = [
        _PROJECT_ROOT / "workspace" / clean,
        _PROJECT_ROOT / "workspace" / "input" / clean,
        _PROJECT_ROOT / "workspace" / "documents" / clean,
        _PROJECT_ROOT / "workspace" / "projects" / clean,
        _PROJECT_ROOT / "workspace" / "output" / clean,
        _PROJECT_ROOT / "data" / "uploads" / clean,
        _PROJECT_ROOT / clean,
    ]
    for c in candidates:
        if c.exists() and c.is_file():
            return c
    return None


def edit_pdf_document(
    file_path: str,
    output_filename: Optional[str] = None,
    watermark_text: Optional[str] = None,
    header_text: Optional[str] = None,
    footer_text: Optional[str] = None,
    append_text: Optional[str] = None,
    output_dir: str = WORKSPACE_OUTPUT_DIR,
) -> dict[str, Any]:
    """
    Read and edit an existing PDF document in the workspace.
    Supports adding watermarks, header/footer annotations, and appending new pages.
    """
    resolved = resolve_workspace_file_path(file_path)
    if not resolved:
        return {
            "success": False,
            "error": f"PDF file '{file_path}' not found in workspace.",
        }

    os.makedirs(output_dir, exist_ok=True)
    if not output_filename:
        base = resolved.stem
        output_filename = f"{base}_edited_{int(time.time())}.pdf"
    elif not output_filename.endswith(".pdf"):
        output_filename = f"{output_filename}.pdf"

    target_path = os.path.join(output_dir, output_filename)

    try:
        import pymupdf as fitz
        doc = fitz.open(str(resolved))

        for page in doc:
            rect = page.rect
            if header_text:
                page.insert_text((50, 30), header_text, fontsize=9, color=(0.2, 0.3, 0.4))
            if footer_text:
                page.insert_text((50, rect.height - 30), footer_text, fontsize=9, color=(0.3, 0.3, 0.3))
            if watermark_text:
                center_x = max(50, rect.width / 2 - (len(watermark_text) * 5))
                center_y = rect.height / 2
                page.insert_text((center_x, center_y), watermark_text, fontsize=20, color=(0.8, 0.2, 0.2))

        if append_text:
            new_page = doc.new_page()
            rect = new_page.rect
            text_rect = fitz.Rect(50, 50, rect.width - 50, rect.height - 50)
            new_page.insert_textbox(text_rect, append_text, fontsize=11, color=(0.1, 0.1, 0.1))

        doc.save(target_path)
        doc.close()
        size_bytes = os.path.getsize(target_path)

        return {
            "success": True,
            "filename": output_filename,
            "file_path": os.path.abspath(target_path),
            "relative_path": f"output/{output_filename}",
            "size_bytes": size_bytes,
            "message": f"Successfully edited PDF document: saved to 'output/{output_filename}' ({size_bytes} bytes).",
        }
    except Exception as e:
        logger.error(f"Failed to edit PDF document: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"PDF edit error: {str(e)}",
        }


def read_workspace_document(
    file_path: str,
    max_chars: int = 8000,
) -> dict[str, Any]:
    """
    Read and extract text from an input document or file in the workspace.
    Supports PDF, DOCX, XLSX, TXT, CSV, Markdown, and JSON files.
    """
    resolved = resolve_workspace_file_path(file_path)
    if not resolved:
        return {
            "success": False,
            "error": f"File '{file_path}' not found in workspace or uploads.",
        }

    ext = resolved.suffix.lower()
    content_text = ""
    try:
        if ext == ".pdf":
            import pymupdf as fitz
            doc = fitz.open(str(resolved))
            pages_text = []
            for i, page in enumerate(doc):
                txt = page.get_text()
                if txt.strip():
                    pages_text.append(f"--- Page {i + 1} ---\n{txt}")
            doc.close()
            content_text = "\n\n".join(pages_text) if pages_text else "(No extractable text found; document may require OCR)"
        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(str(resolved), data_only=True)
            sheets_text = []
            for name in wb.sheetnames:
                ws = wb[name]
                rows = list(ws.iter_rows(values_only=True))
                sheet_lines = [f"### Sheet: {name}"]
                for r in rows[:100]:
                    if any(cell is not None for cell in r):
                        sheet_lines.append(" | ".join(str(c) if c is not None else "" for c in r))
                sheets_text.append("\n".join(sheet_lines))
            wb.close()
            content_text = "\n\n".join(sheets_text)
        elif ext == ".docx":
            import docx
            doc = docx.Document(str(resolved))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            content_text = "\n".join(paras)
        else:
            content_text = resolved.read_text(encoding="utf-8", errors="replace")

        truncated = False
        if len(content_text) > max_chars:
            content_text = content_text[:max_chars] + f"\n...[Truncated: {len(content_text)} total characters]"
            truncated = True

        return {
            "success": True,
            "filename": resolved.name,
            "path": str(resolved.relative_to(_PROJECT_ROOT) if resolved.is_relative_to(_PROJECT_ROOT) else resolved),
            "extension": ext,
            "content": content_text,
            "total_chars": len(content_text),
            "truncated": truncated,
        }
    except Exception as e:
        logger.error(f"Error reading document '{file_path}': {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Error reading document: {str(e)}",
        }


def create_markdown_document(
    title: str,
    content: str,
    output_filename: Optional[str] = None,
    output_dir: str = WORKSPACE_OUTPUT_DIR,
) -> dict[str, Any]:
    """Create a formatted Markdown (.md) document in the local workspace."""
    os.makedirs(output_dir, exist_ok=True)
    if not output_filename:
        clean_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).rstrip()
        filename_base = clean_title.lower().replace(" ", "_") or "document"
        output_filename = f"{filename_base}_{int(time.time())}.md"
    elif not output_filename.endswith(".md"):
        output_filename = f"{output_filename}.md"

    file_path = os.path.join(output_dir, output_filename)

    try:
        with open(file_path, "w", encoding="utf-8") as f:
            if not content.startswith("#"):
                f.write(f"# {title}\n\n{content}\n")
            else:
                f.write(f"{content}\n")

        size_bytes = os.path.getsize(file_path)
        return {
            "success": True,
            "filename": output_filename,
            "file_path": os.path.abspath(file_path),
            "relative_path": f"output/{output_filename}",
            "size_bytes": size_bytes,
            "message": f"Successfully created Markdown document: {output_filename}",
        }
    except Exception as e:
        logger.error(f"Failed to create Markdown document: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Markdown creation error: {str(e)}",
        }


def create_excel_spreadsheet(
    title: str,
    headers: list[str],
    rows: list[list[Any]],
    sheet_name: str = "Sheet1",
    output_filename: Optional[str] = None,
    output_dir: str = WORKSPACE_OUTPUT_DIR,
) -> dict[str, Any]:
    """Create a styled Excel spreadsheet (.xlsx) with headers and data rows."""
    os.makedirs(output_dir, exist_ok=True)
    if not output_filename:
        clean_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).rstrip()
        filename_base = clean_title.lower().replace(" ", "_") or "data_table"
        output_filename = f"{filename_base}_{int(time.time())}.xlsx"
    elif not output_filename.endswith(".xlsx"):
        output_filename = f"{output_filename}.xlsx"

    file_path = os.path.join(output_dir, output_filename)

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name[:31]

        header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin", color="CBD5E1"),
            right=Side(style="thin", color="CBD5E1"),
            top=Side(style="thin", color="CBD5E1"),
            bottom=Side(style="thin", color="CBD5E1"),
        )
        center_align = Alignment(horizontal="center", vertical="center")
        left_align = Alignment(horizontal="left", vertical="center")

        # Write headers
        ws.append(headers)
        for col_num in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border

        # Write data rows
        for row in rows:
            ws.append(row)

        for row_num in range(2, len(rows) + 2):
            for col_num in range(1, len(headers) + 1):
                cell = ws.cell(row=row_num, column=col_num)
                cell.border = thin_border
                if isinstance(cell.value, (int, float)):
                    cell.alignment = Alignment(horizontal="right", vertical="center")
                else:
                    cell.alignment = left_align

        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            col_letter = get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

        wb.save(file_path)
        size_bytes = os.path.getsize(file_path)

        return {
            "success": True,
            "filename": output_filename,
            "file_path": os.path.abspath(file_path),
            "relative_path": f"output/{output_filename}",
            "row_count": len(rows),
            "column_count": len(headers),
            "size_bytes": size_bytes,
            "message": f"Successfully created Excel spreadsheet '{output_filename}' with {len(rows)} rows.",
        }
    except Exception as e:
        logger.error(f"Failed to create Excel spreadsheet: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Excel creation error: {str(e)}",
        }


def edit_excel_spreadsheet(
    file_path: str,
    sheet_name: Optional[str] = None,
    cell_updates: Optional[dict[str, Any]] = None,
    append_rows: Optional[list[list[Any]]] = None,
) -> dict[str, Any]:
    """
    Read and edit an existing Excel spreadsheet (.xlsx) in the workspace.
    Supports updating specific cells (e.g. {'B2': 150.0, 'C3': 'Approved'}) and appending new rows.
    """
    resolved = resolve_workspace_file_path(file_path)
    if not resolved:
        return {
            "success": False,
            "error": f"Excel file '{file_path}' not found in workspace.",
        }
    target_path = str(resolved)

    try:
        import openpyxl
        from openpyxl.styles import Border, Side

        wb = openpyxl.load_workbook(target_path)
        if sheet_name and sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            ws = wb.active

        updated_cells_count = 0
        if cell_updates:
            for cell_ref, val in cell_updates.items():
                ws[cell_ref] = val
                updated_cells_count += 1

        appended_rows_count = 0
        if append_rows:
            thin_border = Border(
                left=Side(style="thin", color="CBD5E1"),
                right=Side(style="thin", color="CBD5E1"),
                top=Side(style="thin", color="CBD5E1"),
                bottom=Side(style="thin", color="CBD5E1"),
            )
            for row in append_rows:
                ws.append(row)
                appended_rows_count += 1
                curr_row = ws.max_row
                for c in range(1, len(row) + 1):
                    cell = ws.cell(row=curr_row, column=c)
                    cell.border = thin_border

        wb.save(target_path)
        size_bytes = os.path.getsize(target_path)

        return {
            "success": True,
            "file_path": os.path.abspath(target_path),
            "filename": os.path.basename(target_path),
            "updated_cells": updated_cells_count,
            "updated_cells_count": updated_cells_count,
            "appended_rows": appended_rows_count,
            "appended_rows_count": appended_rows_count,
            "total_rows": ws.max_row,
            "size_bytes": size_bytes,
            "message": f"Successfully updated Excel spreadsheet '{os.path.basename(target_path)}' ({updated_cells_count} cells modified, {appended_rows_count} rows appended).",
        }
    except Exception as e:
        logger.error(f"Failed to edit Excel spreadsheet: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Excel editing error: {str(e)}",
        }


def generate_report_file(
    doc_type: str,
    title: str,
    content: str,
    output_dir: str = WORKSPACE_OUTPUT_DIR,
    table_data: Optional[list[list[Any]]] = None,
) -> dict[str, Any]:
    """Generate downloadable Word/PDF/Markdown/Excel report artifact."""
    doc_type_clean = doc_type.lower().replace(".", "").strip()

    if doc_type_clean == "pdf":
        return create_pdf_document(title=title, content=content, output_dir=output_dir, table_data=table_data)
    elif doc_type_clean in ("xlsx", "excel"):
        headers = ["Item", "Description", "Value", "Status"]
        rows = [["1", line[:40], "100.00", "Verified"] for line in content.split("\n") if line.strip()][:10]
        if not rows:
            rows = [["1", title, "1.0", "Complete"]]
        return create_excel_spreadsheet(title=title, headers=headers, rows=rows, output_dir=output_dir)
    elif doc_type_clean in ("md", "markdown"):
        return create_markdown_document(title=title, content=content, output_dir=output_dir)
    elif doc_type_clean == "docx":
        os.makedirs(output_dir, exist_ok=True)
        clean_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-")).rstrip()
        filename = f"{clean_title.lower().replace(' ', '_')}_{int(time.time())}.docx"
        file_path = os.path.join(output_dir, filename)
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(title, 0)
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(file_path)
            return {
                "success": True,
                "title": title,
                "doc_type": "docx",
                "filename": filename,
                "file_path": os.path.abspath(file_path),
                "relative_path": f"output/{filename}",
                "size_bytes": os.path.getsize(file_path),
                "message": f"Successfully generated Word document: {filename}",
            }
        except Exception as e:
            return create_markdown_document(title=title, content=content, output_dir=output_dir)
    else:
        return create_markdown_document(title=title, content=content, output_dir=output_dir)


CREATE_PDF_DOCUMENT_TOOL = ToolDefinition(
    name="create_pdf_document",
    description="Generate a professional formatted PDF document or report with headings, text, and optional tables.",
    required_skill="Artifact Drafting",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title of the PDF document."},
            "content": {"type": "string", "description": "Body text or markdown formatted content."},
            "output_filename": {"type": "string", "description": "Optional custom filename (e.g. 'audit_report.pdf')."},
            "table_data": {
                "type": "array",
                "items": {"type": "array", "items": {"type": "string"}},
                "description": "Optional 2D array of rows for a structured data table.",
            },
        },
        "required": ["title", "content"],
    },
    func=create_pdf_document,
)

CREATE_MARKDOWN_DOCUMENT_TOOL = ToolDefinition(
    name="create_markdown_document",
    description="Create a formatted Markdown (.md) note, report, or specification in the workspace.",
    required_skill="Artifact Drafting",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title of the markdown document."},
            "content": {"type": "string", "description": "Markdown body content with headings, lists, tables."},
            "output_filename": {"type": "string", "description": "Optional filename (e.g. 'notes.md')."},
        },
        "required": ["title", "content"],
    },
    func=create_markdown_document,
)

CREATE_EXCEL_SPREADSHEET_TOOL = ToolDefinition(
    name="create_excel_spreadsheet",
    description="Create a styled Microsoft Excel spreadsheet (.xlsx) with custom column headers, rows, and auto-adjusted layout.",
    required_skill="Data Engineering",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Title / topic of the spreadsheet."},
            "headers": {"type": "array", "items": {"type": "string"}, "description": "Column header names."},
            "rows": {
                "type": "array",
                "items": {"type": "array", "items": {"type": "string"}},
                "description": "2D array of row values corresponding to column headers.",
            },
            "sheet_name": {"type": "string", "description": "Worksheet tab name (default: 'Sheet1')."},
            "output_filename": {"type": "string", "description": "Optional filename (e.g. 'procurement.xlsx')."},
        },
        "required": ["title", "headers", "rows"],
    },
    func=create_excel_spreadsheet,
)

EDIT_EXCEL_SPREADSHEET_TOOL = ToolDefinition(
    name="edit_excel_spreadsheet",
    description="Edit an existing Excel (.xlsx) file: update specific cell coordinates (e.g. {'B2': 500, 'D4': 'Approved'}) or append new rows.",
    required_skill="Data Engineering",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Filename or path of the existing Excel file to edit."},
            "cell_updates": {
                "type": "object",
                "description": "Key-value mapping of cell references to new values (e.g. {'B2': 1200, 'C3': 'Completed'}).",
            },
            "append_rows": {
                "type": "array",
                "items": {"type": "array", "items": {"type": "string"}},
                "description": "New data rows to append to the bottom of the spreadsheet.",
            },
        },
        "required": ["file_path"],
    },
    func=edit_excel_spreadsheet,
)


# ==========================================
# 5. Tool Registry & Schemas
# ==========================================

SEARCH_DOCUMENTS_TOOL = ToolDefinition(
    name="search_documents",
    description="Vector similarity search in pgvector database for relevant document chunks and citations.",
    required_skill="Information Retrieval",
    parameters={
        "type": "object",
        "properties": {
            "query": {
                "type": "string",
                "description": "The search query string to find matching context.",
            },
            "top_k": {
                "type": "integer",
                "description": "Number of top matching chunks to retrieve (default: 3).",
                "default": 3,
            },
            "department": {
                "type": "string",
                "description": "Optional department filter (e.g. 'operations', 'compliance').",
            },
        },
        "required": ["query"],
    },
    func=search_documents,
)

GET_DOCUMENT_PAGE_TOOL = ToolDefinition(
    name="get_document_page",
    description="Retrieve raw text and markdown content for a specific document page.",
    required_skill="Precision Reading",
    parameters={
        "type": "object",
        "properties": {
            "doc_id": {
                "type": "string",
                "description": "Unique identifier of the document.",
            },
            "page_number": {
                "type": "integer",
                "description": "Page number to retrieve (1-indexed).",
            },
        },
        "required": ["doc_id", "page_number"],
    },
    func=get_document_page,
)

CALCULATE_EXPRESSION_TOOL = ToolDefinition(
    name="calculate_expression",
    description="Safely evaluate mathematical and quantitative expressions.",
    required_skill="Quantitative Calculation",
    parameters={
        "type": "object",
        "properties": {
            "expression": {
                "type": "string",
                "description": "Mathematical expression (e.g. '150 * 0.18 + sqrt(64)').",
            },
        },
        "required": ["expression"],
    },
    func=calculate_expression,
)

RUN_PYTHON_SANDBOX_TOOL = ToolDefinition(
    name="run_python_sandbox",
    description="Run isolated Python code in a safe sandbox for data analysis, complex algorithms, or transformations.",
    required_skill="Code Execution",
    parameters={
        "type": "object",
        "properties": {
            "code": {
                "type": "string",
                "description": "Python source code to execute.",
            },
            "timeout_seconds": {
                "type": "integer",
                "description": "Execution timeout in seconds (default: 5).",
                "default": 5,
            },
        },
        "required": ["code"],
    },
    func=run_python_sandbox,
)

GENERATE_REPORT_FILE_TOOL = ToolDefinition(
    name="generate_report_file",
    description="Generate downloadable Word, PDF, Markdown, or text report artifacts.",
    required_skill="Artifact Drafting",
    parameters={
        "type": "object",
        "properties": {
            "doc_type": {
                "type": "string",
                "description": "Document extension / format: 'pdf', 'docx', 'markdown', 'txt', 'json'.",
            },
            "title": {
                "type": "string",
                "description": "Title of the report.",
            },
            "content": {
                "type": "string",
                "description": "Content body in markdown format.",
            },
        },
        "required": ["doc_type", "title", "content"],
    },
    func=generate_report_file,
)


EDIT_PDF_DOCUMENT_TOOL = ToolDefinition(
    name="edit_pdf_document",
    description="Edit an existing PDF document in the workspace: overlay watermarks, add header/footer text, or append new content pages.",
    required_skill="Document Engineering",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Filename or path of the existing PDF file to edit."},
            "output_filename": {"type": "string", "description": "Optional output PDF filename."},
            "watermark_text": {"type": "string", "description": "Watermark or stamp text across pages (e.g. 'CONFIDENTIAL', 'APPROVED')."},
            "header_text": {"type": "string", "description": "Header text to place at the top of each page."},
            "footer_text": {"type": "string", "description": "Footer text to place at the bottom of each page."},
            "append_text": {"type": "string", "description": "Text notes or markdown content to append as a new page."},
        },
        "required": ["file_path"],
    },
    func=edit_pdf_document,
)

READ_WORKSPACE_DOCUMENT_TOOL = ToolDefinition(
    name="read_workspace_document",
    description="Read and inspect the text and tabular content of any input document (PDF, Excel XLSX, Word DOCX, TXT, CSV, Markdown, JSON) in the workspace.",
    required_skill="Precision Reading",
    parameters={
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Filename or path of the document to read."},
            "max_chars": {"type": "integer", "description": "Maximum characters to extract (default: 8000).", "default": 8000},
        },
        "required": ["file_path"],
    },
    func=read_workspace_document,
)


class ToolRegistry:
    """Catalog managing all agent tools and execution dispatcher."""

    def __init__(self, load_defaults: bool = True):
        self._tools: dict[str, ToolDefinition] = {}
        if load_defaults:
            self._load_defaults()

    def _load_defaults(self) -> None:
        defaults = [
            SEARCH_DOCUMENTS_TOOL,
            GET_DOCUMENT_PAGE_TOOL,
            CALCULATE_EXPRESSION_TOOL,
            RUN_PYTHON_SANDBOX_TOOL,
            GENERATE_REPORT_FILE_TOOL,
            CREATE_PDF_DOCUMENT_TOOL,
            EDIT_PDF_DOCUMENT_TOOL,
            READ_WORKSPACE_DOCUMENT_TOOL,
            CREATE_MARKDOWN_DOCUMENT_TOOL,
            CREATE_EXCEL_SPREADSHEET_TOOL,
            EDIT_EXCEL_SPREADSHEET_TOOL,
        ]
        for tool in defaults:
            self.register(tool)

    def register(self, tool: ToolDefinition) -> None:
        """Register a new tool definition."""
        self._tools[tool.name] = tool

    def unregister(self, tool_name: str) -> bool:
        """Unregister a tool by name."""
        if tool_name in self._tools:
            del self._tools[tool_name]
            return True
        return False

    def get(self, tool_name: str) -> Optional[ToolDefinition]:
        """Look up tool by name."""
        return self._tools.get(tool_name)

    def get_or_raise(self, tool_name: str) -> ToolDefinition:
        """Get tool definition or raise KeyError."""
        tool = self.get(tool_name)
        if not tool:
            raise KeyError(f"Tool '{tool_name}' not found. Available tools: {list(self._tools.keys())}")
        return tool

    def list_tools(self) -> list[ToolDefinition]:
        """List all registered tools."""
        return list(self._tools.values())

    def get_ollama_tools(self, tool_names: Optional[list[str]] = None) -> list[dict[str, Any]]:
        """Get list of tool schemas formatted for Ollama / OpenAI API."""
        if tool_names is None:
            tools = self._tools.values()
        else:
            tools = [self._tools[name] for name in tool_names if name in self._tools]
        return [t.to_ollama_tool() for t in tools]

    def execute(self, tool_name: str, arguments: dict[str, Any]) -> dict[str, Any]:
        """Execute a tool synchronously with structured error handling."""
        tool = self.get(tool_name)
        if not tool:
            return {
                "success": False,
                "error": f"Tool '{tool_name}' is not registered in ToolRegistry.",
            }

        try:
            if inspect.iscoroutinefunction(tool.func):
                return asyncio.run(tool.func(**arguments))
            return tool.func(**arguments)
        except Exception as e:
            logger.error(f"Error executing tool '{tool_name}' with args {arguments}: {e}", exc_info=True)
            return {
                "success": False,
                "error": f"Execution failed for tool '{tool_name}': {str(e)}",
            }

    async def aexecute(self, tool_name: str, arguments: dict[str, Any]) -> dict[str, Any]:
        """Execute a tool asynchronously."""
        tool = self.get(tool_name)
        if not tool:
            return {
                "success": False,
                "error": f"Tool '{tool_name}' is not registered in ToolRegistry.",
            }

        try:
            if inspect.iscoroutinefunction(tool.func):
                return await tool.func(**arguments)
            return await asyncio.to_thread(tool.func, **arguments)
        except Exception as e:
            logger.error(f"Error executing tool '{tool_name}' with args {arguments}: {e}", exc_info=True)
            return {
                "success": False,
                "error": f"Execution failed for tool '{tool_name}': {str(e)}",
            }


# Global default tool registry instance
tool_registry = ToolRegistry()
