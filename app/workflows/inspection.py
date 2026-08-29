from pathlib import Path

from docx import Document

from app.rag.service import RAGService


class InspectionWorkflow:
    """Agentic workflow for inspection-report analysis."""

    def __init__(self, service: RAGService | None = None) -> None:
        self.service = service or RAGService()

    def run(
        self,
        document_path: Path,
        output_path: Path,
    ) -> dict:
        """
        Ingest one inspection report, analyze only its indexed
        chunks, and produce a DOCX summary.
        """

        document_path = Path(document_path)
        output_path = Path(output_path)

        if not document_path.exists():
            raise FileNotFoundError(
                f"Inspection document not found: {document_path}"
            )

        # Agent step 1: ingest the document.
        ingestion = self.service.ingest(document_path)

        # Agent step 2: retrieve only this document's chunks.
        contexts = self.service.vector_store.get_by_source(
            document_path.name,
        )

        if not contexts:
            raise RuntimeError(
                "No indexed context was found for the "
                f"document: {document_path.name}"
            )

        # Agent step 3: ask the configured generator to analyze
        # only the retrieved document context.
        answer = self.service.generator.generate(
            (
                "Analyze this inspection report and provide a "
                "concise inspection summary. Identify the "
                "document ID, equipment tag, equipment, plant, "
                "inspection findings, operating conditions, and "
                "any other important information explicitly "
                "present in the report."
            ),
            contexts,
        )

        # Agent step 4: produce a DOCX deliverable.
        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document()

        document.add_heading(
            "Inspection Summary",
            level=1,
        )

        document.add_paragraph(
            f"Source: {document_path.name}"
        )

        document.add_paragraph(
            f"Chunks analyzed: {len(contexts)}"
        )

        document.add_heading(
            "Findings and Assessment",
            level=2,
        )

        document.add_paragraph(answer)

        document.add_heading(
            "Approval Note",
            level=2,
        )

        document.add_paragraph(
            "This summary was generated from the indexed "
            "inspection report. The responsible maintenance "
            "team should review the findings before approval "
            "or further action."
        )

        document.save(str(output_path))

        return {
            "source": document_path.name,
            "chunks_indexed": ingestion["chunks_indexed"],
            "chunks_analyzed": len(contexts),
            "vectors_stored": ingestion["vectors_stored"],
            "answer": answer,
            "output": str(output_path),
        }
